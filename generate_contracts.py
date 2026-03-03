"""
=============================================================
TOOL SINH BIÊN BẢN NGHIỆM THU & THANH LÝ  v3
=============================================================
Cấu trúc thư mục:
  template_nghiem_thu.docx   ← file Word mẫu BBNT
  template_thanh_ly.docx     ← file Word mẫu BBTL
  danh_sach.xlsx             ← mỗi sheet = 1 học viên
  generate_bienban.py        ← file này

Chạy: python generate_bienban.py
Cài: pip install python-docx openpyxl
=============================================================

PLACEHOLDER trong Word mẫu:
  Thông tin học viên / bảo lãnh:
    {{TEN_HV}}  {{NGAY_SINH}}  {{CCCD_HV}}  {{NGAY_CAP_HV}}
    {{NOI_CAP_HV}}  {{DIA_CHI}}  {{SDT_HV}}
    {{TEN_BL}}  {{CCCD_BL}}  {{NGAY_CAP_BL}}  {{NOI_CAP_BL}}
    {{SDT_BL}}  {{QUAN_HE}}

  Số hợp đồng:
    {{SO_HOP_DONG}}   → số HĐ gốc  vd: 220901/2025/DHD2-TOPIK3456-TBT
    {{SO_BBNT}}       → tự sinh:   220901/2025/BBNT-TBT  (lấy prefix + /BBNT-TBT)
    {{SO_BBTL}}       → tự sinh:   220901/2025/BBTL-TBT
    {{SO_BBNT_PREFIX}}→ phần đầu số BBNT cũ (nếu Word tách run, không cần nhập)
    {{SO_BBNT_SUFFIX}}→ /BBNT-TBT  (không cần nhập)
    {{NGAY_KY_HD}}    → ngày ký hợp đồng gốc

  Thanh toán (file thanh lý):
    {{GD1_TIEN}}  {{GD2_TIEN}}  {{GD3_TIEN}}  ...  (tối đa 5 giai đoạn)
    {{TONG_TIEN}}   → tự tính từ các giai đoạn

CÁCH NHẬP EXCEL (mỗi sheet = 1 học viên):
  Cột A: Label          |  Cột B: Giá trị
  Số hợp đồng           |  220901/2025/DHD2-TOPIK3456-TBT
  Ngày ký hợp đồng      |  22/09/2025
  ...thông tin học viên...
  Giai đoạn 1           |  28.080.000   ← đặt CUỐI CÙNG
  Giai đoạn 2           |  38.400.000
  Giai đoạn 3           |  9.876.000
=============================================================
"""

import os
import re
import copy
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import openpyxl

# ============================================================
# CẤU HÌNH
# ============================================================
TEMPLATE_NGHIEM_THU = "template_nghiem_thu.docx"
TEMPLATE_THANH_LY   = "template_thanh_ly.docx"
EXCEL_FILE          = "danh_sach.xlsx"
OUTPUT_DIR          = "output"

LABEL_MAP = {
    # Số hợp đồng
    "số hợp đồng":                  "SO_HOP_DONG",
    "ngày ký hợp đồng":             "NGAY_KY_HD",
    # Học viên
    "họ và tên":                    "TEN_HV",
    "ngày sinh":                    "NGAY_SINH",
    "cccd số":                      "CCCD_HV",
    "ccdc số":                      "CCCD_HV",
    "ngày cấp cccd":                "NGAY_CAP_HV",
    "ngày cấp":                     "NGAY_CAP_HV",
    "nơi cấp cccd":                 "NOI_CAP_HV",
    "nơi cấp":                      "NOI_CAP_HV",
    "địa chỉ":                      "DIA_CHI",
    "số điện thoại":                "SDT_HV",
    # Bảo lãnh
    "họ và tên (bảo lãnh)":         "TEN_BL",
    "cccd số (bảo lãnh)":           "CCCD_BL",
    "ccdc số (bảo lãnh)":           "CCCD_BL",
    "ngày cấp cccd (bảo lãnh)":     "NGAY_CAP_BL",
    "ngày cấp (bảo lãnh)":          "NGAY_CAP_BL",
    "nơi cấp cccd (bảo lãnh)":      "NOI_CAP_BL",
    "nơi cấp (bảo lãnh)":           "NOI_CAP_BL",
    "số điện thoại (bảo lãnh)":     "SDT_BL",
    "quan hệ với học viên":         "QUAN_HE",
    # Giai đoạn thanh toán  (label "Giai đoạn 1", "Giai đoạn 2", ...)
    "giai đoạn 1":                  "GD1_TIEN",
    "giai đoạn 2":                  "GD2_TIEN",
    "giai đoạn 3":                  "GD3_TIEN",
    "giai đoạn 4":                  "GD4_TIEN",
    "giai đoạn 5":                  "GD5_TIEN",
    # Ngày thanh toán
    "ngày thanh toán 1":             "GD1_NGAY",
    "ngày thanh toán 2":             "GD2_NGAY",
    "ngày thanh toán 3":             "GD3_NGAY",
    "ngày thanh toán 4":             "GD4_NGAY",
    "ngày thanh toán 5":             "GD5_NGAY",
}
# ============================================================


# ──────────────────────────────────────────────
#  ĐỌC EXCEL
# ──────────────────────────────────────────────

def parse_number(s):
    """'28.080.000' → 28080000. Trả về None nếu không parse được."""
    if not s:
        return None
    clean = str(s).replace(".", "").replace(",", "").replace(" ", "").replace("VNĐ", "").replace("vnđ", "")
    try:
        return int(clean)
    except ValueError:
        return None


def format_number(n):
    """28080000 → '28.080.000'"""
    return f"{n:,.0f}".replace(",", ".")


def read_excel_sheets(excel_file):
    wb = openpyxl.load_workbook(excel_file, data_only=False)
    students = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        info = {"_sheet": sheet_name, "_giai_doan": []}
        for row in ws.iter_rows(values_only=False):
            if not row[0] or not row[0].value:
                continue
            label = str(row[0].value).strip()
            
            # Get formatted value (what user sees in Excel) or raw value
            cell = row[1] if len(row) > 1 else None
            if cell and cell.value is not None:
                # Try to get formatted value first
                value = cell.value
                if hasattr(cell, 'number_format') and cell.number_format and cell.number_format != 'General':
                    # Cell has formatting - try to use the formatted display value
                    try:
                        from openpyxl.styles.numbers import FORMAT_DATE_XLSX14, FORMAT_DATE_DATETIME
                        from datetime import datetime
                        if isinstance(value, datetime):
                            value = value.strftime("%d/%m/%Y")
                        else:
                            value = str(value)
                    except:
                        value = str(value)
                else:
                    value = str(value).strip()
            else:
                value = ""
            
            key = LABEL_MAP.get(label.lower())
            if key:
                info[key] = value
                # Lưu riêng danh sách giai đoạn theo thứ tự
                if key.startswith("GD") and key.endswith("_TIEN"):
                    n = parse_number(value)
                    if n is not None:
                        gd_num = int(key[2])  # GD1 → 1
                        info["_giai_doan"].append((gd_num, n))

        # Sắp xếp giai đoạn theo số thứ tự
        info["_giai_doan"].sort(key=lambda x: x[0])

        # Tính tổng và điền lại placeholder
        if info["_giai_doan"]:
            tong = sum(v for _, v in info["_giai_doan"])
            info["TONG_TIEN"] = format_number(tong)
            info["SO_GIAI_DOAN"] = str(len(info["_giai_doan"]))
            for gd_num, val in info["_giai_doan"]:
                info[f"GD{gd_num}_TIEN"] = format_number(val)

        # Tự sinh số biên bản từ SO_HOP_DONG
        # VD: "220901/2025/DHD2-TOPIK3456-TBT" → prefix = "220901/2025"
        so_hd = info.get("SO_HOP_DONG", "")
        if so_hd:
            prefix = so_hd.split("/")[0] + "/" + so_hd.split("/")[1] if so_hd.count("/") >= 1 else so_hd
            info.setdefault("SO_BBNT", f"{prefix}/BBNT-TBT")
            info.setdefault("SO_BBTL", f"{prefix}/BBTL-TBT")
            info.setdefault("SO_BBNT_PREFIX", prefix)
            info.setdefault("SO_BBNT_SUFFIX", "/BBNT-TBT")

        students.append(info)
    return students


# ──────────────────────────────────────────────
#  THAY PLACEHOLDER TRONG PARAGRAPH
# ──────────────────────────────────────────────

def replace_in_paragraph(para, replacements):
    full = "".join(r.text for r in para.runs)
    new = full
    for key, val in replacements.items():
        new = new.replace(f"{{{{{key}}}}}", str(val) if val else "")
    if new == full or not para.runs:
        return
    para.runs[0].text = new
    for r in para.runs[1:]:
        r.text = ""


def replace_all(doc, replacements):
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, replacements)


# ──────────────────────────────────────────────
#  XỬ LÝ LIST GIAI ĐOẠN ĐỘNG (file Nghiệm Thu)
# ──────────────────────────────────────────────

def handle_payment_list(doc, giai_doan_list):
    """
    Tìm các paragraph chứa {{DONG_GIAI_DOAN}}, {{DONG_GD_2}}, {{DONG_GD_3}}
    trong file nghiệm thu, xây lại đúng số dòng theo giai_doan_list.
    """
    if not giai_doan_list:
        return

    # Tìm các paragraph placeholder giai đoạn
    GD_PLACEHOLDERS = {"{{DONG_GIAI_DOAN}}", "{{DONG_GD_2}}", "{{DONG_GD_3}}",
                       "{{DONG_GD_4}}", "{{DONG_GD_5}}"}
    gd_paras = []
    for para in doc.paragraphs:
        txt = "".join(r.text for r in para.runs)
        if txt.strip() in GD_PLACEHOLDERS:
            gd_paras.append(para)

    if not gd_paras:
        return

    # Paragraph mẫu để clone (lấy cái đầu tiên)
    template_para = gd_paras[0]
    parent = template_para._p.getparent()

    # Vị trí chèn = vị trí paragraph đầu tiên
    insert_idx = list(parent).index(template_para._p)

    # Xoá tất cả paragraph placeholder cũ
    for para in gd_paras:
        parent.remove(para._p)

    # Tạo paragraph mới cho từng giai đoạn, chèn theo thứ tự
    for i, (gd_num, amount) in enumerate(giai_doan_list):
        new_p = copy.deepcopy(template_para._p)
        # Tìm w:t trong new_p và set text
        for t_elem in new_p.iter(qn("w:t")):
            t_elem.text = f"+ Giai đoạn {gd_num} : {format_number(amount)}"
            t_elem.attrib.pop("{http://www.w3.org/XML/1998/namespace}space", None)
        parent.insert(insert_idx + i, new_p)


# ──────────────────────────────────────────────
#  XỬ LÝ BẢNG GIAI ĐOẠN ĐỘNG (chỉ file Thanh Lý)
# ──────────────────────────────────────────────

def handle_payment_table(doc, giai_doan_list, info):
    """
    Template có sẵn rows với placeholders {{GD1_TIEN}}, {{GD1_NGAY}}, etc.
    Xóa các rows không dùng (ví dụ nếu chỉ có 2 giai đoạn, xóa rows GD3, GD4, GD5).
    Placeholders sẽ được replace bởi replace_all() sau này.
    """
    if not giai_doan_list:
        return

    payment_table = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Giai đoạn" in cell.text and "Số tiền" in "".join(
                    c.text for r in table.rows for c in r.cells
                ):
                    payment_table = table
                    break
            if payment_table:
                break
        if payment_table:
            break

    if payment_table is None:
        return

    tbl_elem = payment_table._tbl
    all_tr = tbl_elem.findall(qn("w:tr"))
    
    if len(all_tr) < 3:
        return
    
    header_tr = all_tr[0]
    total_tr = all_tr[-1]
    data_rows = all_tr[1:-1]
    
    num_gd = len(giai_doan_list)
    
    # Remove rows for unused giai doan (keep only first num_gd data rows)
    for i in range(num_gd, len(data_rows)):
        tbl_elem.remove(data_rows[i])


# ──────────────────────────────────────────────
#  FILL DOC
# ──────────────────────────────────────────────

def fill_doc(template_path, info, is_thanh_ly=False):
    doc = Document(template_path)
    replacements = {k: v for k, v in info.items() if not k.startswith("_")}

    if is_thanh_ly and info.get("_giai_doan"):
        handle_payment_table(doc, info["_giai_doan"], info)
    elif not is_thanh_ly and info.get("_giai_doan"):
        handle_payment_list(doc, info["_giai_doan"])

    replace_all(doc, replacements)
    return doc


def sanitize(name):
    return re.sub(r'[\\/:*?"<>|]', '_', str(name)).strip()


# ──────────────────────────────────────────────
#  MAIN
# ──────────────────────────────────────────────

def main():
    print("\n" + "="*55)
    print("  SINH BIÊN BẢN NGHIỆM THU & THANH LÝ  v3")
    print("="*55)

    missing = [f for f in [TEMPLATE_NGHIEM_THU, TEMPLATE_THANH_LY, EXCEL_FILE]
               if not os.path.exists(f)]
    if missing:
        for f in missing: print(f"❌ Không tìm thấy: {f}")
        return

    Path(OUTPUT_DIR).mkdir(exist_ok=True)
    students = read_excel_sheets(EXCEL_FILE)
    print(f"📊 Đọc được {len(students)} học viên\n")

    ok = err = 0
    for info in students:
        ten  = info.get("TEN_HV", info["_sheet"])
        name = sanitize(ten)
        gd   = info.get("_giai_doan", [])
        tong = info.get("TONG_TIEN", "?")
        print(f"  🔄 {ten}  |  {len(gd)} giai đoạn  |  Tổng: {tong}")

        try:
            doc_nt = fill_doc(TEMPLATE_NGHIEM_THU, info, is_thanh_ly=False)
            doc_nt.save(os.path.join(OUTPUT_DIR, f"BBNT_{name}.docx"))

            doc_tl = fill_doc(TEMPLATE_THANH_LY, info, is_thanh_ly=True)
            doc_tl.save(os.path.join(OUTPUT_DIR, f"BBTL_{name}.docx"))

            print(f"    ✅ BBNT_{name}.docx")
            print(f"    ✅ BBTL_{name}.docx")
            ok += 1
        except Exception as e:
            import traceback
            print(f"    ❌ Lỗi: {e}")
            traceback.print_exc()
            err += 1

    print(f"\n{'='*55}")
    print(f"  Xong: {ok} học viên ({ok*2} file), {err} lỗi")
    print(f"  Kết quả: {os.path.abspath(OUTPUT_DIR)}/")
    print("="*55 + "\n")


if __name__ == "__main__":
    main()